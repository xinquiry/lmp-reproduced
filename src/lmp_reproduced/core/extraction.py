"""
Module for extracting content from Word documents.

This module provides functions to extract text, tables, and image counts
from Word (.docx) documents.
"""

from pathlib import Path
from typing import Optional

from docx import Document

from lmp_reproduced.models import DocumentContent, Paragraph, Table


def extract_document(docx_path: Path | str) -> DocumentContent:
    """
    Extract all content from a Word document.

    Args:
        docx_path: Path to the Word document

    Returns:
        DocumentContent containing paragraphs, tables, and image count
    """
    docx_path = Path(docx_path)
    doc = Document(docx_path)

    # Extract paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(Paragraph(
                text=para.text.strip(),
                style=para.style.name if para.style else None
            ))

    # Extract tables
    tables = []
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
        tables.append(Table(index=table_idx, rows=rows))

    # Count embedded objects/images (approximate)
    images_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images_count += 1

    return DocumentContent(
        paragraphs=paragraphs,
        tables=tables,
        images_count=images_count,
        source_file=docx_path
    )


def save_extracted_content(
    content: DocumentContent,
    output_file: Path | str
) -> Path:
    """
    Save extracted content to a text file.

    Args:
        content: DocumentContent to save
        output_file: Path to output text file

    Returns:
        Path to the written file
    """
    output_file = Path(output_file)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    lines = []
    lines.append(f"FILE: {content.source_file.name}")
    lines.append("=" * 80)
    lines.append("")

    lines.append("--- Document Statistics ---")
    lines.append(f"Paragraphs: {len(content.paragraphs)}")
    lines.append(f"Tables: {len(content.tables)}")
    lines.append(f"Images: {content.images_count}")
    lines.append("")

    lines.append("--- Full Content ---")
    lines.append("")
    for para in content.paragraphs:
        style_info = f"[{para.style}]" if para.style else "[None]"
        lines.append(f"{style_info}: {para.text}")

    lines.append("")
    lines.append("--- Tables ---")
    for table in content.tables:
        lines.append(f"\nTable {table.index}:")
        for row in table.rows:
            lines.append(f"  | {' | '.join(row)} |")

    output_file.write_text("\n".join(lines), encoding="utf-8")
    return output_file


def extract_all_documents(
    reports_dir: Path | str,
    output_dir: Path | str
) -> list[DocumentContent]:
    """
    Extract content from all Word documents in a directory.

    Args:
        reports_dir: Directory containing Word documents
        output_dir: Directory to save extracted content

    Returns:
        List of DocumentContent for each processed document
    """
    reports_dir = Path(reports_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(exist_ok=True, parents=True)

    results = []

    for docx_file in reports_dir.glob("*.docx"):
        content = extract_document(docx_file)

        # Generate safe filename
        safe_name = docx_file.stem.replace(" ", "_").replace("—", "-")
        output_file = output_dir / f"{safe_name}.txt"

        save_extracted_content(content, output_file)
        results.append(content)

    return results
